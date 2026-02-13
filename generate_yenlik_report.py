#!/usr/bin/env python3
# -*- coding: utf-8 -*-

import pandas as pd
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from datetime import datetime
import os

def set_table_borders(table):
    """Add borders to table"""
    tbl = table._tbl
    tblPr = tbl.tblPr
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    
    # Create table borders element
    tblBorders = OxmlElement('w:tblBorders')
    
    # Define border style
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')  # Border size
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')  # Black color
        tblBorders.append(border)
    
    tblPr.append(tblBorders)

def parse_track_data(track_string):
    """Parse track data from the format: 'Track Name (€amount, streams стримов)'"""
    tracks = []
    for track in track_string.split(' | '):
        try:
            name_part = track.split(' (€')[0].strip()
            data_part = track.split('(€')[1].split(')')[0]
            parts = data_part.split(', ')
            
            revenue_str = parts[0].strip()
            streams_str = parts[1].strip().replace(' стримов', '').replace(',', '')
            
            revenue = float(revenue_str)
            streams = int(streams_str)
            
            tracks.append({
                'name': name_part,
                'revenue': revenue,
                'streams': streams
            })
        except:
            continue
    
    return tracks

def get_platform_stats_from_csv(csv_files, artist_name='Yenlik'):
    """Get platform statistics for an artist from CSV files"""
    platform_stats = {}
    
    print(f"\n🔍 Поиск данных по артисту '{artist_name}' в CSV файлах...")
    
    for csv_file in csv_files:
        if not os.path.exists(csv_file):
            print(f"⚠️  Файл не найден: {csv_file}")
            continue
            
        try:
            print(f"\n📂 Обработка: {csv_file}")
            df = pd.read_csv(csv_file, sep=';', encoding='utf-8', low_memory=False)
            
            # Filter for artist (including collaborations like "Yenlik, rauana")
            yenlik_mask = df['Исполнитель'].str.contains(artist_name, na=False, case=False)
            artist_data = df[yenlik_mask].copy()
            
            print(f"   Найдено записей: {len(artist_data)}")
            
            if len(artist_data) == 0:
                continue
            
            # Convert revenue from string (European format with comma) to float
            artist_data['Сумма вознаграждения'] = (
                artist_data['Сумма вознаграждения']
                .astype(str)
                .str.replace(',', '.')
                .astype(float)
            )
            
            # Group by platform
            platform_groups = artist_data.groupby('Платформа').agg({
                'Количество': 'sum',
                'Сумма вознаграждения': 'sum'
            }).reset_index()
            
            # Add to total stats
            for _, row in platform_groups.iterrows():
                platform = row['Платформа']
                streams = int(row['Количество'])
                revenue = float(row['Сумма вознаграждения'])
                
                if platform in platform_stats:
                    platform_stats[platform]['streams'] += streams
                    platform_stats[platform]['revenue'] += revenue
                else:
                    platform_stats[platform] = {
                        'streams': streams,
                        'revenue': revenue
                    }
            
        except Exception as e:
            print(f"❌ Ошибка при обработке {csv_file}: {e}")
            continue
    
    print(f"\n✅ Найдено уникальных платформ: {len(platform_stats)}")
    
    return platform_stats

def create_yenlik_report():
    """Create report for Yenlik using template"""
    
    print("=" * 70)
    print("🎵 Генерация отчета для артиста Yenlik")
    print("=" * 70)
    
    # Read artist data from analysis
    df = pd.read_csv('full_artists_analysis.csv')
    yenlik_data = df[df['Артист'] == 'Yenlik'].iloc[0]
    
    # Parse track data
    tracks = parse_track_data(yenlik_data['Все треки (с доходами)'])
    top_tracks = sorted(tracks, key=lambda x: x['revenue'], reverse=True)[:5]
    
    # Calculate totals
    total_revenue = float(yenlik_data['Общий доход (EUR)'])
    total_streams = int(yenlik_data['Всего стримов'])
    num_tracks = int(yenlik_data['Всего треков'])
    
    # Get platform stats
    csv_files = [
        '1740260_704133_2025-07-01_2025-09-01 2.csv',
        '1855874_704133_2025-10-01_2025-12-01 (1).csv'
    ]
    
    platform_stats = get_platform_stats_from_csv(csv_files)
    top_platforms = sorted(platform_stats.items(), key=lambda x: x[1]['revenue'], reverse=True)[:5]
    
    # Load original template
    print("\n📄 Загрузка шаблона...")
    doc = Document('ozen_template_final.docx')
    
    # Replace placeholders in the first page (keep original formatting)
    print("📝 Замена плейсхолдеров...")
    
    # Prepare replacement values
    distribution_payment = total_revenue * 0.75
    copyright_payment = 0.00
    total_payment = distribution_payment + copyright_payment
    
    # Create list of releases (track names)
    releases_list = ', '.join([t['name'] for t in tracks[:5]]) + f' и другие ({num_tracks} треков всего)'
    
    replacements = {
        '{{ARTIST_NAME}}': 'Yenlik',
        '{{RELEASES}}': releases_list,
        '{{PERIOD}}': 'Июль - Декабрь 2025',
        '{{DISTRIBUTION_PAYMENT}': f'€{distribution_payment:,.2f}',  # Note: template has typo without closing }
        '{{DISTRIBUTION_PAYMENT}}': f'€{distribution_payment:,.2f}',  # Also handle correct version
        '{{COPYRIGHT_PAYMENT}}': f'€{copyright_payment:,.2f}',
        '{{TOTAL_PAYMENT}}': f'€{total_payment:,.2f}'
    }
    
    # Replace in all paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                # Replace while preserving formatting
                for run in paragraph.runs:
                    if key in run.text:
                        run.text = run.text.replace(key, value)
    
    # ========== PAGE 2: Analytics ==========
    print("📊 Добавление страницы с аналитикой...")
    
    doc.add_page_break()
    
    # Title
    title_para = doc.add_paragraph('Аналитика по артисту Yenlik')
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.runs[0]
    title_run.font.size = Pt(18)
    title_run.font.bold = True
    
    doc.add_paragraph()
    
    # General stats
    stats_heading = doc.add_paragraph('Общая статистика')
    stats_heading.runs[0].font.size = Pt(14)
    stats_heading.runs[0].font.bold = True
    
    stats_para = doc.add_paragraph()
    stats_para.add_run(f'Всего треков: ').bold = True
    stats_para.add_run(f'{num_tracks}\n')
    
    stats_para.add_run(f'Всего стримов: ').bold = True
    stats_para.add_run(f'{total_streams:,}')
    
    doc.add_paragraph()
    
    # Top 5 tracks
    tracks_heading = doc.add_paragraph('Топ-5 треков по доходу')
    tracks_heading.runs[0].font.size = Pt(14)
    tracks_heading.runs[0].font.bold = True
    
    table = doc.add_table(rows=1, cols=4)
    # Don't set style - template doesn't have standard styles
    set_table_borders(table)  # Add borders
    
    # Header
    header_cells = table.rows[0].cells
    header_cells[0].text = '№'
    header_cells[1].text = 'Трек'
    header_cells[2].text = 'Стримы'
    header_cells[3].text = '% от общего дохода'
    
    for cell in header_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
    
    # Add tracks
    for idx, track in enumerate(top_tracks, 1):
        row_cells = table.add_row().cells
        row_cells[0].text = str(idx)
        row_cells[1].text = track['name']
        row_cells[2].text = f"{track['streams']:,}"
        percentage = (track['revenue'] / total_revenue) * 100
        row_cells[3].text = f"{percentage:.1f}%"
    
    doc.add_paragraph()
    
    # Top 5 platforms
    platforms_heading = doc.add_paragraph('Топ-5 платформ по доходу')
    platforms_heading.runs[0].font.size = Pt(14)
    platforms_heading.runs[0].font.bold = True
    
    if top_platforms:
        platform_table = doc.add_table(rows=1, cols=4)
        # Don't set style - template doesn't have standard styles
        set_table_borders(platform_table)  # Add borders
        
        # Header
        header_cells = platform_table.rows[0].cells
        header_cells[0].text = '№'
        header_cells[1].text = 'Платформа'
        header_cells[2].text = 'Стримы'
        header_cells[3].text = '% от общего дохода'
        
        for cell in header_cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
        
        # Add platforms
        for idx, (platform, stats) in enumerate(top_platforms, 1):
            row_cells = platform_table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[1].text = platform
            row_cells[2].text = f"{stats['streams']:,}"
            percentage = (stats['revenue'] / total_revenue) * 100
            row_cells[3].text = f"{percentage:.1f}%"
    else:
        doc.add_paragraph('⚠️ Данные по платформам не найдены')
    
    doc.add_paragraph()
    
    # Footer
    footer_para = doc.add_paragraph()
    footer_para.add_run('\n\nОтчет сгенерирован автоматически\n').italic = True
    footer_para.add_run(f'Дата создания: {datetime.now().strftime("%d.%m.%Y %H:%M")}').italic = True
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save
    output_file = 'yenlik_report.docx'
    doc.save(output_file)
    
    # Print summary
    print("\n" + "=" * 70)
    print(f"✅ Отчет успешно создан: {output_file}")
    print("=" * 70)
    
    print(f"\n📊 Статистика:")
    print(f"   Всего треков: {num_tracks}")
    print(f"   Общий доход: €{total_revenue:,.2f}")
    print(f"   Всего стримов: {total_streams:,}")
    print(f"   Выплата артисту (75%): €{total_payment:,.2f}")
    
    print(f"\n🎵 Топ-5 треков:")
    for idx, track in enumerate(top_tracks, 1):
        percentage = (track['revenue'] / total_revenue) * 100
        print(f"   {idx}. {track['name']}")
        print(f"      {track['streams']:,} стримов | {percentage:.1f}% от дохода")
    
    if top_platforms:
        print(f"\n📱 Топ-5 платформ:")
        for idx, (platform, stats) in enumerate(top_platforms, 1):
            percentage = (stats['revenue'] / total_revenue) * 100
            print(f"   {idx}. {platform}")
            print(f"      {stats['streams']:,} стримов | {percentage:.1f}% от дохода")

if __name__ == '__main__':
    create_yenlik_report()
